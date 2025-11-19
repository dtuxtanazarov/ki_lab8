import requests
from bs4 import BeautifulSoup
import time
from docx import Document

BASE_URL = "https://lex.uz"
START_URL = BASE_URL + "/uz/"
OUTPUT_DOCX = "lex_all_texts.docx"

visited = set()
to_visit = set([START_URL])
all_texts = []

def get_links(url):
    try:
        r = requests.get(url, timeout=10)
        if r.status_code != 200:
            return set(), []
        soup = BeautifulSoup(r.text, "html.parser")

        # Extract visible text
        texts = [tag.get_text(strip=True) for tag in soup.find_all(["p", "div", "span"])]
        texts = [t for t in texts if t]  

        links = set()
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if href.startswith("/"):
                href = BASE_URL + href
            if href.startswith(BASE_URL):
                links.add(href)

        return links, texts
    except Exception as e:
        print("❌ Error fetching", url, ":", e)
        return set(), []

# Crawl loop
while to_visit and len(visited) < 50:
    url = to_visit.pop()
    if url in visited:
        continue
    visited.add(url)

    print(f"🔎 Crawling: {url}")
    links, texts = get_links(url)

    if texts:
        all_texts.extend(texts)
        print(f"✅ Extracted {len(texts)} text blocks")

    # Add new links to visit
    new_links = links - visited
    to_visit.update(new_links)

    time.sleep(1)  # be polite

# Save results to Word
doc = Document()
doc.add_heading("Lex.uz Scraped Data", level=1)

for t in all_texts:
    doc.add_paragraph(t)

doc.save(OUTPUT_DOCX)
print(f"\n📄 Done! Saved {len(all_texts)} text blocks from {len(visited)} pages into {OUTPUT_DOCX}")
